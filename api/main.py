from fastapi import FastAPI, Form
import pandas as pd
import difflib
import requests
import io
from docx import Document
from fastapi.responses import FileResponse
import uuid

app = FastAPI()

def classify_discrepancy(gt_text, ocr_line, score):
    gt_upper = str(gt_text).upper()
    ocr_upper = str(ocr_line).upper()

    if any(sym in gt_text for sym in ["¹", "²"]) and ("?" in ocr_line or "?" in gt_text):
        discrepancy_type = "Symbol mismatch / line-structure difference"
        details = "Footnote marker ¹ appears as ? in image OCR; placement differs."
    elif score < 0.95 and ("." in ocr_line or "-" in ocr_line or "…" in ocr_line):
        discrepancy_type = "Order / punctuation difference"
        details = "Number appears after descriptor; dot-leaders/extra characters present."
    elif score < 0.8:
        discrepancy_type = "Wording difference"
        details = "Text wording differs significantly."
    else:
        discrepancy_type = "Minor formatting difference"
        details = "Formatting or spacing mismatch."

    if "NITROGEN" in gt_upper or "NUTRIENT" in gt_upper:
        category = "Nutrient declaration"
    elif "INGREDIENT" in gt_upper or "CMC" in gt_upper:
        category = "Ingredients / footnote"
    elif "GRANULOMETRY" in gt_upper or "PRILLS" in gt_upper:
        category = "Granulometry"
    elif "STORAGE" in gt_upper or "WARNING" in gt_upper:
        category = "Safety / storage"
    else:
        category = "General"

    severity = "Critical" if category in ["Nutrient declaration", "Ingredients / footnote"] else "Low"
    return discrepancy_type, details, category, severity


def save_report_doc(report_rows):
    """Generate DOCX report with a unique filename."""
    filename = f"discrepancy_report_{uuid.uuid4().hex}.docx"
    doc = Document()
    doc.add_heading("Discrepancy Report", level=1)

    table = doc.add_table(rows=1, cols=8)
    hdr_cells = table.rows[0].cells
    headers = ["Language", "#", "Discrepancy Type", "Discrepancy Details",
               "Excel Text (Approved)", "Image Text (Printed)", "Category", "Severity"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for row in report_rows:
        row_cells = table.add_row().cells
        row_cells[0].text = row["Language"]
        row_cells[1].text = str(row["#"])
        row_cells[2].text = row["Discrepancy Type"]
        row_cells[3].text = row["Discrepancy Details"]
        row_cells[4].text = row["Excel Text (Approved)"]
        row_cells[5].text = row["Image Text (Printed)"]
        row_cells[6].text = row["Category"]
        row_cells[7].text = row["Severity"]

    doc.save(filename)
    return filename


@app.post("/compare")
async def compare_ocr(csv_url: str = Form(...), ocr_text: str = Form(...)):
    try:
        response = requests.get(csv_url)
        if response.status_code != 200:
            return {"error": f"Failed to fetch file from {csv_url}. Status: {response.status_code}"}

        file_bytes = io.BytesIO(response.content)

        # Try Excel first
        try:
            df = pd.read_excel(file_bytes, header=0, engine="openpyxl")
        except Exception as e1:
            file_bytes.seek(0)
            try:
                df = pd.read_csv(file_bytes, header=0)
            except Exception as e2:
                return {"error": f"Could not parse file. Excel error: {str(e1)}, CSV error: {str(e2)}"}

        if 0 in df.index:
            df = df.drop(index=0)

        ocr_lines = [line.strip() for line in ocr_text.splitlines() if line.strip()]

        report_rows = []
        discrepancy_id = 1

        for lang in df.columns:
            for gt_text in df[lang].dropna():
                for ocr_line in ocr_lines:
                    score = difflib.SequenceMatcher(None, ocr_line, str(gt_text)).ratio()
                    if score < 0.95:
                        discrepancy_type, details, category, severity = classify_discrepancy(gt_text, ocr_line, score)
                        report_rows.append({
                            "Language": lang,
                            "#": f"{lang} {discrepancy_id}.",
                            "Discrepancy Type": discrepancy_type,
                            "Discrepancy Details": details,
                            "Excel Text (Approved)": str(gt_text),
                            "Image Text (Printed)": str(ocr_line),
                            "Category": category,
                            "Severity": severity
                        })
                        discrepancy_id += 1

        report_file = save_report_doc(report_rows)

        return FileResponse(
            report_file,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=report_file
        )

    except Exception as e:
        return {"error": f"Unexpected error: {str(e)}"}


@app.get("/")
async def root():
    return {"message": "OCR Compare API is running. Use POST /compare."}
